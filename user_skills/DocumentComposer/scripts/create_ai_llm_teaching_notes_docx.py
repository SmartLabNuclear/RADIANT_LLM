"""
Create a Microsoft Word teacher-notes companion document for the PowerPoint deck:
"AI & Large Language Models: A visual teaching deck for high school students".

Outputs:
    AI_LLMs_Generative_AI_Teaching_Notes.docx

Dependencies:
    pip install python-docx

Usage:
    python create_ai_llm_teaching_notes_docx.py

The document is designed as teaching material for high school classes. It includes:
- Course framing and learning objectives
- Suggested pacing
- Slide-by-slide teacher notes for a 21-slide deck
- Discussion prompts, checks for understanding, and activities
- Glossary, assessment ideas, and responsible-use guidance
"""

from __future__ import annotations

from pathlib import Path
from typing import Iterable

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_CELL_VERTICAL_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor
from docx.enum.style import WD_STYLE_TYPE


OUTPUT = Path(__file__).resolve().parent / "AI_LLMs_Generative_AI_Teaching_Notes.docx"

NAVY = RGBColor(18, 32, 58)
BLUE = RGBColor(43, 111, 246)
GREEN = RGBColor(25, 170, 112)
PURPLE = RGBColor(124, 77, 255)
ORANGE = RGBColor(245, 153, 61)
RED = RGBColor(226, 80, 86)
MID = RGBColor(91, 110, 140)
LIGHT_FILL = "E9F0FF"
PALE_GREEN = "E8F8F1"
PALE_ORANGE = "FFF4E6"
PALE_PURPLE = "F2ECFF"
WHITE = "FFFFFF"
BORDER = "D9E2F2"


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = tc_pr.find(qn("w:shd"))
    if shd is None:
        shd = OxmlElement("w:shd")
        tc_pr.append(shd)
    shd.set(qn("w:fill"), fill)


def set_cell_border(cell, color: str = BORDER, size: str = "8") -> None:
    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    borders = tc_pr.first_child_found_in("w:tcBorders")
    if borders is None:
        borders = OxmlElement("w:tcBorders")
        tc_pr.append(borders)
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        tag = "w:" + edge
        element = borders.find(qn(tag))
        if element is None:
            element = OxmlElement(tag)
            borders.append(element)
        element.set(qn("w:val"), "single")
        element.set(qn("w:sz"), size)
        element.set(qn("w:space"), "0")
        element.set(qn("w:color"), color)


def set_repeat_table_header(row) -> None:
    tr_pr = row._tr.get_or_add_trPr()
    tbl_header = OxmlElement("w:tblHeader")
    tbl_header.set(qn("w:val"), "true")
    tr_pr.append(tbl_header)


def add_page_number(paragraph) -> None:
    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = paragraph.add_run()
    fld_char1 = OxmlElement("w:fldChar")
    fld_char1.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = "PAGE"
    fld_char2 = OxmlElement("w:fldChar")
    fld_char2.set(qn("w:fldCharType"), "end")
    run._r.append(fld_char1)
    run._r.append(instr)
    run._r.append(fld_char2)


def add_hyperlink(paragraph, text: str, url: str) -> None:
    # Add an external hyperlink using python-docx OOXML support.
    part = paragraph.part
    r_id = part.relate_to(
        url,
        "http://schemas.openxmlformats.org/officeDocument/2006/relationships/hyperlink",
        is_external=True,
    )
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    r_pr = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), "2B6FF6")
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    r_pr.append(color)
    r_pr.append(underline)
    run.append(r_pr)
    t = OxmlElement("w:t")
    t.text = text
    run.append(t)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def add_colored_heading(doc: Document, text: str, level: int = 1, color: RGBColor = NAVY):
    p = doc.add_heading(text, level=level)
    for run in p.runs:
        run.font.color.rgb = color
        run.font.name = "Aptos Display"
    return p


def add_callout(doc: Document, title: str, body: str, fill: str = LIGHT_FILL, border: str = "2B6FF6") -> None:
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = table.cell(0, 0)
    set_cell_shading(cell, fill)
    set_cell_border(cell, color=border, size="10")
    cell.vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.CENTER
    p = cell.paragraphs[0]
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(title)
    r.bold = True
    r.font.color.rgb = NAVY
    r.font.size = Pt(11)
    p2 = cell.add_paragraph(body)
    p2.paragraph_format.space_after = Pt(0)
    for run in p2.runs:
        run.font.size = Pt(10)
    doc.add_paragraph()


def add_bullets(doc: Document, items: Iterable[str], style: str = "List Bullet") -> None:
    for item in items:
        p = doc.add_paragraph(item, style=style)
        p.paragraph_format.space_after = Pt(2)


def add_numbered(doc: Document, items: Iterable[str]) -> None:
    for item in items:
        p = doc.add_paragraph(item, style="List Number")
        p.paragraph_format.space_after = Pt(2)


def add_table(doc: Document, headers: list[str], rows: list[list[str]], widths: list[float] | None = None) -> None:
    table = doc.add_table(rows=1, cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    hdr = table.rows[0]
    set_repeat_table_header(hdr)
    for i, header in enumerate(headers):
        cell = hdr.cells[i]
        cell.text = header
        set_cell_shading(cell, LIGHT_FILL)
        set_cell_border(cell)
        for p in cell.paragraphs:
            for run in p.runs:
                run.bold = True
                run.font.color.rgb = NAVY
    for row in rows:
        cells = table.add_row().cells
        for i, value in enumerate(row):
            cells[i].text = value
            set_cell_border(cells[i])
            cells[i].vertical_alignment = WD_CELL_VERTICAL_ALIGNMENT.TOP
    if widths:
        for row in table.rows:
            for idx, width in enumerate(widths):
                row.cells[idx].width = Inches(width)
    doc.add_paragraph()


SLIDE_NOTES = [
    {
        "num": 1,
        "title": "AI & Large Language Models",
        "objective": "Introduce the unit and set expectations: AI is a tool students can learn to understand, question, and use responsibly.",
        "talk": "Start by asking students where they have already seen AI: search engines, recommendation feeds, translation, image filters, chatbots, games, phones, and school tools. Explain that this lesson is not about hype; it is about understanding how the technology works well enough to use it thoughtfully.",
        "activity": "Think-pair-share: name one AI tool you have used and one question you have about how it works.",
        "misconception": "AI is not a magic brain. It is software built from data, mathematics, computing, and human design choices.",
        "check": "Students can name at least two AI applications and one responsible-use concern.",
    },
    {
        "num": 2,
        "title": "Learning goals",
        "objective": "Preview the path through the lesson: definitions, mechanism, usage, modern app patterns, and safety.",
        "talk": "Walk through each goal quickly. Emphasize that students are not expected to become AI engineers in one class, but they should be able to explain core ideas clearly and evaluate outputs critically.",
        "activity": "Have students mark each goal as: I know this, I have heard of this, or this is new to me.",
        "misconception": "Knowing how to use a chatbot is not the same as understanding AI literacy. The lesson includes both skills and judgment.",
        "check": "Students can identify which learning goal they most want to understand by the end.",
    },
    {
        "num": 3,
        "title": "What is AI?",
        "objective": "Define AI in student-friendly terms as software that performs tasks usually associated with human intelligence.",
        "talk": "Give concrete examples: recognizing a face, recommending a video, predicting the next word, detecting spam, or helping with translation. Then distinguish narrow AI from human general intelligence.",
        "activity": "Sort examples into AI or not AI: calculator, spam filter, autocomplete, thermostat, chess engine, human student.",
        "misconception": "AI systems do not need to be conscious to be useful. Most AI is narrow and task-specific.",
        "check": "Students can explain AI as pattern-based software rather than a human-like mind.",
    },
    {
        "num": 4,
        "title": "AI, ML, Deep Learning, Generative AI",
        "objective": "Clarify the vocabulary stack and show how LLMs fit within broader AI.",
        "talk": "Use nested categories: AI is the broad field; machine learning learns from examples; deep learning uses neural networks with many layers; generative AI creates content; LLMs are generative AI systems focused on language.",
        "activity": "Ask students to place four sticky notes on a board: AI, ML, Deep Learning, LLM. Then arrange them from broadest to most specific.",
        "misconception": "Not all AI is generative AI, and not all generative AI is an LLM.",
        "check": "Students can state one example of each category.",
    },
    {
        "num": 5,
        "title": "What makes AI generative?",
        "objective": "Explain that generative AI creates new content that resembles patterns learned from training examples.",
        "talk": "Compare a classifier and a generator. A classifier might label a picture as a cat. A generator might create a new cat image from a prompt. Discuss text, images, code, audio, and video as output types.",
        "activity": "Students list three outputs a generative AI could create for a science class, an art class, and an English class.",
        "misconception": "Generated content is not automatically original, correct, ethical, or free to use. It must be reviewed.",
        "check": "Students can distinguish creating a new response from choosing a label.",
    },
    {
        "num": 6,
        "title": "Large Language Models: the core idea",
        "objective": "Teach the next-token prediction idea without oversimplifying it into a toy autocomplete.",
        "talk": "Explain tokens as chunks of text. The model looks at previous tokens and estimates likely next tokens. It repeats this process many times to produce sentences, paragraphs, code, or explanations.",
        "activity": "Use the phrase 'The student opened the...' and ask students to predict possible next words. Discuss why some words are more likely in context.",
        "misconception": "An LLM is not simply copying a stored answer. It generates by using learned statistical patterns, although training data can still influence outputs.",
        "check": "Students can describe generation as repeated prediction of likely next tokens.",
    },
    {
        "num": 7,
        "title": "Inside an LLM: tokens, embeddings, attention",
        "objective": "Introduce the main internal stages: tokenization, vectors, attention, prediction, and repetition.",
        "talk": "Tokens are converted into numbers called embeddings. Attention helps the model weigh which earlier words matter for the next prediction. Use a sentence where a pronoun refers to an earlier noun to show why context matters.",
        "activity": "Give the sentence: 'Maria gave Ava her notebook because she forgot hers.' Ask: what does each pronoun refer to, and why is context important?",
        "misconception": "Embeddings are not meanings in a human sense; they are numeric representations that capture useful relationships learned from data.",
        "check": "Students can define attention as a way to connect relevant context across text.",
    },
    {
        "num": 8,
        "title": "How models are trained",
        "objective": "Explain the difference between pretraining, fine-tuning, feedback, and deployment.",
        "talk": "Training is like large-scale practice: the system improves at predicting tokens by comparing predictions to examples. Alignment and feedback shape behavior so the model is more helpful, safe, and instruction-following.",
        "activity": "Analogy discussion: how is model training like practicing a sport or instrument, and how is it different?",
        "misconception": "Training does not guarantee truth. A model can learn style and patterns without having direct access to current facts.",
        "check": "Students can explain why more training data does not automatically mean perfect answers.",
    },
    {
        "num": 9,
        "title": "What can LLMs do well?",
        "objective": "Identify productive uses: tutoring, drafting, summarizing, coding support, brainstorming, and accessibility.",
        "talk": "Frame LLMs as thinking partners rather than answer machines. They are often strongest when helping organize, explain, revise, or generate options. They are weaker when unchecked factual accuracy matters.",
        "activity": "Students rewrite a weak use case into a strong learning use case. Example: 'Write my essay' becomes 'Ask me questions to help plan my essay.'",
        "misconception": "Using AI well does not mean asking it to do all the thinking. The student remains responsible for understanding and final work.",
        "check": "Students can classify a use case as helpful learning support or risky shortcut.",
    },
    {
        "num": 10,
        "title": "Limits: why LLMs can be wrong",
        "objective": "Explain hallucination, bias, stale knowledge, limited memory, and privacy risk.",
        "talk": "Emphasize that fluency is not proof. A model can produce text that sounds confident but is incorrect. Students should verify facts, sources, calculations, and claims, especially for schoolwork or decisions.",
        "activity": "Show two short AI-style answers, one accurate and one containing subtle errors. Students identify what needs checking.",
        "misconception": "A confident tone does not equal reliability. Verification is a normal part of AI use.",
        "check": "Students can name at least three reasons an LLM answer may be unreliable.",
    },
    {
        "num": 11,
        "title": "Prompt Engineering: how to ask better questions",
        "objective": "Teach prompts as instructions with role, task, context, format, constraints, and checking.",
        "talk": "Compare vague and specific prompts. A vague prompt asks 'Explain climate change.' A stronger prompt specifies audience, length, examples, vocabulary level, and a request to note uncertainty.",
        "activity": "Students improve a vague prompt using role, task, context, format, and constraints.",
        "misconception": "Prompt engineering is not tricking the model. It is clear communication and task design.",
        "check": "Students can add missing context or constraints to improve a prompt.",
    },
    {
        "num": 12,
        "title": "Prompt pattern: CLEAR",
        "objective": "Give students a memorable prompt checklist: Context, Limits, Examples, Ask, Review.",
        "talk": "Have students read the example prompt and identify each CLEAR part. Explain that Review means asking the model to check assumptions, uncertainty, or possible mistakes.",
        "activity": "Students write a CLEAR prompt for one current class assignment, then trade with a partner for feedback.",
        "misconception": "A long prompt is not automatically a good prompt. A good prompt is purposeful and specific.",
        "check": "Students can label the five CLEAR elements in their own prompt.",
    },
    {
        "num": 13,
        "title": "RAG: Retrieval-Augmented Generation",
        "objective": "Explain RAG as a pattern where AI retrieves relevant evidence before answering.",
        "talk": "Use a school example: a student asks about the late-work policy. Instead of relying on memory, a RAG system searches the class syllabus and includes the relevant section in the prompt before generating an answer.",
        "activity": "Give students a small paragraph of evidence and ask them to answer a question using only that evidence.",
        "misconception": "RAG does not make answers automatically correct. The retrieved evidence must be relevant and the answer must use it faithfully.",
        "check": "Students can describe the RAG chain: question, search, evidence, answer, sources.",
    },
    {
        "num": 14,
        "title": "RAG pipeline: one practical view",
        "objective": "Show the technical workflow: upload content, chunk and embed, store, retrieve, generate.",
        "talk": "Explain chunking as breaking documents into searchable pieces. Explain embeddings as numeric meaning maps. Explain vector databases as fast similarity search systems. Keep the level conceptual, not implementation-heavy.",
        "activity": "Cut a short article into paragraphs. Students choose which paragraph would be retrieved for three different questions.",
        "misconception": "The model does not read the entire library every time. It usually receives only selected relevant chunks.",
        "check": "Students can explain why citations are useful in RAG outputs.",
    },
    {
        "num": 15,
        "title": "CAG: Cache-Augmented Generation",
        "objective": "Contrast CAG with RAG and explain when cached context is useful.",
        "talk": "Use the library analogy. RAG searches the library at question time. CAG keeps important pages already open or preloaded. CAG can be helpful when the knowledge base is stable and repeatedly used.",
        "activity": "Students decide whether RAG or CAG is better for: today's weather, a class syllabus, a live sports score, a school handbook, and a rapidly changing news story.",
        "misconception": "Caching is not the same as truth. If the cached information is old, the answer may be old too.",
        "check": "Students can name one situation where CAG may be faster or simpler than RAG.",
    },
    {
        "num": 16,
        "title": "MCP: Model Context Protocol",
        "objective": "Introduce MCP as a standardized way to connect AI apps to tools and data sources.",
        "talk": "Describe MCP as a universal adapter. An AI app can connect to approved servers for files, databases, calendars, web APIs, or other tools. Stress permission boundaries and safety controls.",
        "activity": "Students draw an AI assistant connected to three approved tools and write one permission rule for each.",
        "misconception": "Tool access is not unlimited. Good systems restrict what the AI can see or do.",
        "check": "Students can explain why standardized tool connections are useful and why permissions matter.",
    },
    {
        "num": 17,
        "title": "Agents and Agent Skills",
        "objective": "Explain agents as systems that can plan, use tools, check progress, and complete multi-step tasks.",
        "talk": "Contrast a simple chatbot response with an agentic workflow. A chatbot may answer one question. An agent may plan steps, search sources, write code, check output, revise, and deliver a final file. Skills are reusable workflows that make specialized tasks more consistent.",
        "activity": "Students outline an agent workflow for planning a school club event: goal, plan, tools, checks, final output.",
        "misconception": "Agents are not automatically autonomous in a safe way. Human oversight is essential for important actions.",
        "check": "Students can identify at least two tools an agent might use and one check it should perform.",
    },
    {
        "num": 18,
        "title": "Responsible use in school",
        "objective": "Set norms for academic integrity, privacy, verification, and learning-centered AI use.",
        "talk": "Make the rule practical: AI can help students learn, brainstorm, practice, and revise, but submitted work should reflect the student's understanding. Encourage students to ask teachers what AI use is allowed for each assignment.",
        "activity": "Class creates a two-column AI use policy: allowed learning support and not allowed shortcuts.",
        "misconception": "Responsible use is not simply 'never use AI.' It means matching use to the learning goal and being transparent.",
        "check": "Students can explain why copying an AI answer can harm learning even if the answer is correct.",
    },
    {
        "num": 19,
        "title": "Class activity: build a mini RAG answer",
        "objective": "Let students experience evidence-grounded answering and verification.",
        "talk": "Organize students into small groups. Each group chooses a question, finds two reliable passages, asks an AI or writes a response using only the evidence, then marks claims as supported or unsupported.",
        "activity": "Mini RAG lab: question, evidence, prompt, answer, verify, revise. Require students to highlight evidence that supports each major claim.",
        "misconception": "Adding evidence is not enough. Students must check whether the final wording truly follows from the evidence.",
        "check": "Each group submits one revised answer with evidence labels and one unsupported claim they removed or corrected.",
    },
    {
        "num": 20,
        "title": "Key takeaways",
        "objective": "Review the major concepts and connect them to student habits.",
        "talk": "Revisit the phrase 'Think first. Ask well. Check carefully.' Ask students how it applies to prompts, RAG, agents, and responsible use.",
        "activity": "Exit ticket: define one term, give one useful AI prompt strategy, and name one verification step.",
        "misconception": "AI literacy is not a one-time lesson. The tools will change, but the habits of questioning, checking, and ethical use remain important.",
        "check": "Students can summarize the unit in three sentences using at least three key terms.",
    },
    {
        "num": 21,
        "title": "Sources and further reading",
        "objective": "Encourage source-aware teaching and continued learning.",
        "talk": "Explain that AI changes quickly, so teachers and students should use current, reliable references. Discuss the difference between a vendor guide, a research paper, a news article, and a classroom policy.",
        "activity": "Students compare two sources on the same AI topic and judge which is more appropriate for school research.",
        "misconception": "A source is not reliable just because it is linked or cited. Evaluate author, date, evidence, and purpose.",
        "check": "Students can explain why current sources matter for AI topics.",
    },
]


GLOSSARY = [
    ("Artificial Intelligence (AI)", "Software designed to perform tasks that usually require human intelligence, such as recognizing patterns, predicting outcomes, or generating language."),
    ("Machine Learning (ML)", "A type of AI where systems learn patterns from examples instead of only following fixed rules."),
    ("Deep Learning", "Machine learning using multi-layer neural networks that can learn complex patterns from large data sets."),
    ("Generative AI", "AI that creates new content such as text, images, code, audio, or video."),
    ("Large Language Model (LLM)", "A generative AI model trained on language patterns to predict and produce text or code."),
    ("Token", "A chunk of text, such as a word, part of a word, punctuation mark, or symbol, used by language models."),
    ("Embedding", "A numeric representation of text that helps computers compare meaning or similarity."),
    ("Attention", "A mechanism that helps a model focus on relevant parts of the context when making predictions."),
    ("Hallucination", "A confident-sounding AI output that includes false, unsupported, or invented information."),
    ("Prompt Engineering", "The practice of writing clear instructions, context, examples, and constraints to get better AI outputs."),
    ("RAG", "Retrieval-Augmented Generation: a system pattern where the AI retrieves relevant evidence before generating an answer."),
    ("CAG", "Cache-Augmented Generation: a pattern where stable information is preloaded or cached so the model can use it efficiently."),
    ("MCP", "Model Context Protocol: a standard approach for connecting AI apps to approved tools and data sources."),
    ("Agent", "An AI system that can plan steps, use tools, check results, and work toward a goal."),
    ("Agent Skill", "A reusable workflow, instruction set, or tool package that helps an agent perform a specialized task consistently."),
]


def setup_document() -> Document:
    doc = Document()
    section = doc.sections[0]
    section.top_margin = Inches(0.65)
    section.bottom_margin = Inches(0.65)
    section.left_margin = Inches(0.75)
    section.right_margin = Inches(0.75)

    styles = doc.styles
    styles["Normal"].font.name = "Aptos"
    styles["Normal"].font.size = Pt(10.5)
    styles["Normal"].font.color.rgb = RGBColor(32, 45, 71)

    for style_name, size, color in [
        ("Title", 28, NAVY),
        ("Heading 1", 18, NAVY),
        ("Heading 2", 14, BLUE),
        ("Heading 3", 12, GREEN),
    ]:
        style = styles[style_name]
        style.font.name = "Aptos Display"
        style.font.size = Pt(size)
        style.font.color.rgb = color
        style.font.bold = True

    if "Teacher Note Label" not in styles:
        label = styles.add_style("Teacher Note Label", WD_STYLE_TYPE.CHARACTER)
        label.font.bold = True
        label.font.color.rgb = NAVY

    header = section.header.paragraphs[0]
    header.text = "AI & LLMs for High School - Teacher Notes"
    header.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    for run in header.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = MID

    footer = section.footer.paragraphs[0]
    add_page_number(footer)
    for run in footer.runs:
        run.font.size = Pt(9)
        run.font.color.rgb = MID

    return doc


def add_cover(doc: Document) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("AI & Large Language Models")
    r.bold = True
    r.font.name = "Aptos Display"
    r.font.size = Pt(30)
    r.font.color.rgb = NAVY

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Teacher Notes and Classroom Guide")
    r.font.name = "Aptos Display"
    r.font.size = Pt(18)
    r.font.color.rgb = BLUE

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run("Companion document for the high school teaching presentation on Generative AI, LLMs, RAG, CAG, Prompt Engineering, MCP, and Agent Skills")
    r.font.size = Pt(11)
    r.font.color.rgb = MID

    add_callout(
        doc,
        "How to use this guide",
        "Use the slide-by-slide notes as a flexible script. The document includes discussion prompts, activities, misconceptions to address, and checks for understanding. Adapt the examples to your school policy and subject area.",
        fill=LIGHT_FILL,
        border="2B6FF6",
    )

    add_table(
        doc,
        ["Audience", "Suggested length", "Materials", "Best use"],
        [[
            "High school students, grades 9-12",
            "One 60-90 minute lesson, or two 45-minute class periods",
            "Presentation deck, projector, student devices or printed handouts, short evidence passages for the RAG activity",
            "AI literacy lesson, digital citizenship unit, computer science introduction, research skills lesson, or advisory seminar",
        ]],
        widths=[1.7, 1.7, 2.3, 2.1],
    )

    doc.add_page_break()


def add_overview(doc: Document) -> None:
    add_colored_heading(doc, "1. Lesson overview", 1)
    doc.add_paragraph(
        "This lesson introduces students to AI and modern large language model applications. "
        "The goal is balanced AI literacy: students should understand the basic mechanism, know useful classroom applications, and recognize limits and safety issues."
    )

    add_colored_heading(doc, "Essential questions", 2)
    add_bullets(doc, [
        "What is AI, and how are AI, machine learning, deep learning, generative AI, and LLMs related?",
        "How can an LLM generate useful text if it is predicting tokens?",
        "How do prompt engineering, RAG, CAG, MCP, and agents improve or extend basic chatbot behavior?",
        "What does responsible AI use look like in school?",
    ])

    add_colored_heading(doc, "Learning objectives", 2)
    add_bullets(doc, [
        "Define AI, Generative AI, LLMs, Prompt Engineering, RAG, CAG, MCP, Agents, and Agent Skills in student-friendly language.",
        "Describe an LLM generation loop using tokens, embeddings, attention, and next-token prediction.",
        "Write a clear prompt using role, task, context, format, constraints, and review steps.",
        "Compare RAG and CAG and identify when each is useful.",
        "Explain why tool access through MCP and agent workflows require permissions, guardrails, and human oversight.",
        "Apply responsible-use habits: protect privacy, verify facts, cite sources, and use AI to support learning rather than replace it.",
    ])

    add_colored_heading(doc, "Suggested pacing", 2)
    add_table(
        doc,
        ["Time", "Slides", "Focus", "Teacher move"],
        [
            ["5-10 min", "1-4", "Opening, vocabulary stack", "Ask students where they already encounter AI."],
            ["15-20 min", "5-10", "Generative AI and LLM mechanics", "Use token prediction and attention examples."],
            ["10-15 min", "11-12", "Prompt engineering", "Have students revise weak prompts."],
            ["15-20 min", "13-17", "Modern AI application patterns", "Compare RAG, CAG, MCP, and agents using school examples."],
            ["10-20 min", "18-21", "Responsible use and activity", "Run the mini RAG activity or assign it as a group task."],
        ],
        widths=[1.0, 0.9, 2.2, 3.5],
    )

    add_callout(
        doc,
        "Teacher framing",
        "Avoid presenting AI as either a miracle or a disaster. Present it as a powerful tool with strengths, limits, human design choices, and social consequences.",
        fill=PALE_GREEN,
        border="19AA70",
    )


def add_slide_notes(doc: Document) -> None:
    add_colored_heading(doc, "2. Slide-by-slide teaching notes", 1)
    doc.add_paragraph("Each slide includes a teaching objective, suggested talking points, an activity or prompt, a misconception to address, and a quick check for understanding.")

    for note in SLIDE_NOTES:
        add_colored_heading(doc, f"Slide {note['num']}: {note['title']}", 2)
        table = doc.add_table(rows=5, cols=2)
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.CENTER
        labels = [
            ("Teaching objective", note["objective"], LIGHT_FILL),
            ("Suggested teacher script", note["talk"], WHITE),
            ("Discussion or activity", note["activity"], PALE_GREEN),
            ("Common misconception", note["misconception"], PALE_ORANGE),
            ("Check for understanding", note["check"], PALE_PURPLE),
        ]
        for row, (label, text, fill) in zip(table.rows, labels):
            row.cells[0].text = label
            row.cells[1].text = text
            set_cell_shading(row.cells[0], fill)
            set_cell_border(row.cells[0])
            set_cell_border(row.cells[1])
            row.cells[0].width = Inches(1.7)
            row.cells[1].width = Inches(5.8)
            for p in row.cells[0].paragraphs:
                for run in p.runs:
                    run.bold = True
                    run.font.color.rgb = NAVY
        doc.add_paragraph()


def add_activities(doc: Document) -> None:
    add_colored_heading(doc, "3. Classroom activities", 1)

    add_colored_heading(doc, "Activity A: Prompt makeover", 2)
    doc.add_paragraph("Goal: Students learn how context and constraints improve AI outputs.")
    add_numbered(doc, [
        "Give students a weak prompt, such as: Explain World War II.",
        "Ask them to rewrite it using role, task, context, format, constraints, and a review request.",
        "Students trade prompts and identify the strongest and weakest parts.",
        "Class discusses how prompt quality affects answer quality.",
    ])
    add_callout(doc, "Sample stronger prompt", "Act as a patient 10th grade history tutor. Explain three major causes of World War II in under 200 words. Use simple language, include one analogy, and end with two review questions. Note any oversimplifications.", fill=LIGHT_FILL, border="2B6FF6")

    add_colored_heading(doc, "Activity B: Mini RAG lab", 2)
    doc.add_paragraph("Goal: Students practice evidence-grounded answering.")
    add_numbered(doc, [
        "Give each group a short source packet or ask them to choose two reliable passages.",
        "Students write a question that can be answered from the passages.",
        "Students draft an answer using only the evidence.",
        "They underline every claim and mark it supported, unsupported, or needs clarification.",
        "They revise the answer and add source notes.",
    ])

    add_colored_heading(doc, "Activity C: Agent workflow map", 2)
    doc.add_paragraph("Goal: Students understand agents as multi-step systems with tools and checks.")
    add_numbered(doc, [
        "Choose a practical goal, such as planning a fundraiser or creating a study guide.",
        "Students list the agent goal, subtasks, tools, data needed, and safety checks.",
        "Students identify at least one action that should require human approval.",
        "Groups present their workflow and explain the guardrails.",
    ])

    add_colored_heading(doc, "Exit ticket", 2)
    add_bullets(doc, [
        "Define one term from the lesson in your own words.",
        "Write one example of a good AI learning use and one risky use.",
        "Name one way to verify an AI answer.",
    ])


def add_assessment_and_policy(doc: Document) -> None:
    add_colored_heading(doc, "4. Assessment ideas", 1)
    add_table(
        doc,
        ["Assessment", "What it measures", "Suggested scoring"],
        [
            ["Vocabulary quick quiz", "Understanding of key terms", "1 point for accurate definition, 1 point for example."],
            ["Prompt engineering task", "Ability to write clear AI instructions", "Check for role, task, context, format, constraints, and review."],
            ["RAG evidence answer", "Evidence use and verification", "Score supported claims, citation quality, and revision quality."],
            ["Responsible-use reflection", "Ethical reasoning", "Look for privacy awareness, transparency, and learning-centered use."],
            ["Concept map", "Relationships among AI concepts", "Students connect AI, ML, deep learning, generative AI, LLMs, RAG, CAG, MCP, and agents."],
        ],
        widths=[1.8, 2.5, 3.2],
    )

    add_colored_heading(doc, "Suggested responsible-use classroom norms", 2)
    add_bullets(doc, [
        "Use AI to ask for hints, feedback, examples, practice questions, outlines, or explanations.",
        "Do not submit AI-generated work as your own unless the assignment explicitly allows that use.",
        "Never paste private information, passwords, personal identifiers, or confidential school data into AI tools.",
        "Verify factual claims with reliable sources, especially for science, history, health, law, finance, or current events.",
        "Be transparent: when required, explain what AI tool was used and how it helped.",
        "Final work should reflect the student's understanding, voice, and reasoning.",
    ])

    add_callout(doc, "Policy reminder", "School and district AI policies vary. Adapt these notes to local rules before presenting them as classroom policy.", fill=PALE_ORANGE, border="F5993D")


def add_glossary(doc: Document) -> None:
    add_colored_heading(doc, "5. Student-friendly glossary", 1)
    rows = [[term, definition] for term, definition in GLOSSARY]
    add_table(doc, ["Term", "Student-friendly definition"], rows, widths=[2.2, 5.4])


def add_references(doc: Document) -> None:
    add_colored_heading(doc, "6. Teacher reference links", 1)
    doc.add_paragraph("Use current official documentation and recent educational resources when updating this lesson. AI tools and terminology change quickly.")
    links = [
        ("OpenAI documentation", "https://platform.openai.com/docs"),
        ("Microsoft Learn: Prompt engineering", "https://learn.microsoft.com/azure/ai-services/openai/concepts/prompt-engineering"),
        ("Model Context Protocol documentation", "https://modelcontextprotocol.io/"),
        ("Anthropic: Model Context Protocol", "https://www.anthropic.com/news/model-context-protocol"),
        ("IBM: Retrieval-Augmented Generation overview", "https://www.ibm.com/topics/retrieval-augmented-generation"),
    ]
    for label, url in links:
        p = doc.add_paragraph(style="List Bullet")
        add_hyperlink(p, label, url)


def build_docx(output: Path = OUTPUT) -> Path:
    doc = setup_document()
    add_cover(doc)
    add_overview(doc)
    add_slide_notes(doc)
    add_activities(doc)
    add_assessment_and_policy(doc)
    add_glossary(doc)
    add_references(doc)
    doc.save(output)
    return output


if __name__ == "__main__":
    path = build_docx(OUTPUT)
    print(f"Saved teaching notes to: {path}")
